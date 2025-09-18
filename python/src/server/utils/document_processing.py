"""
Document Processing Utilities

This module provides utilities for extracting text from various document formats
including PDF, Word documents, and plain text files.
"""

import io

# Removed direct logging import - using unified config

# Import document processing libraries with availability checks
try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config.logfire_config import get_logger, logfire

logger = get_logger(__name__)


def _preserve_code_blocks_across_pages(text: str) -> str:
    """
    Fix code blocks that were split across PDF page boundaries.
    
    PDFs often break markdown code blocks with page headers like:
    ```python
    def hello():
    --- Page 2 ---
        return "world"
    ```
    
    This function rejoins split code blocks by removing page separators
    that appear within code blocks.
    """
    import re
    
    # Pattern to match page separators that split code blocks
    # Look for: ``` [content] --- Page N --- [content] ```
    page_break_in_code_pattern = r'(```\w*[^\n]*\n(?:[^`]|`(?!``))*)(\n--- Page \d+ ---\n)((?:[^`]|`(?!``))*)```'
    
    # Keep merging until no more splits are found
    while True:
        matches = list(re.finditer(page_break_in_code_pattern, text, re.DOTALL))
        if not matches:
            break
            
        # Replace each match by removing the page separator
        for match in reversed(matches):  # Reverse to maintain positions
            before_page_break = match.group(1)
            page_separator = match.group(2) 
            after_page_break = match.group(3)
            
            # Rejoin the code block without the page separator
            rejoined = f"{before_page_break}\n{after_page_break}```"
            text = text[:match.start()] + rejoined + text[match.end():]
    
    return text


def _clean_html_to_text(html_content: str) -> str:
    """
    Clean HTML tags and convert to plain text suitable for RAG.
    Preserves code blocks and important structure while removing markup.
    """
    import re
    
    # First preserve code blocks with their content before general cleaning
    # This ensures code blocks remain intact for extraction
    code_blocks = []
    
    # Find and temporarily replace code blocks to preserve them
    code_patterns = [
        r'<pre><code[^>]*>(.*?)</code></pre>',
        r'<code[^>]*>(.*?)</code>',
        r'<pre[^>]*>(.*?)</pre>',
    ]
    
    processed_html = html_content
    placeholder_map = {}
    
    for pattern in code_patterns:
        matches = list(re.finditer(pattern, processed_html, re.DOTALL | re.IGNORECASE))
        for i, match in enumerate(reversed(matches)):  # Reverse to maintain positions
            # Extract code content and clean HTML entities
            code_content = match.group(1)
            # Clean HTML entities and span tags from code
            code_content = re.sub(r'<span[^>]*>', '', code_content)
            code_content = re.sub(r'</span>', '', code_content)
            code_content = re.sub(r'&lt;', '<', code_content)
            code_content = re.sub(r'&gt;', '>', code_content)
            code_content = re.sub(r'&amp;', '&', code_content)
            code_content = re.sub(r'&quot;', '"', code_content)
            code_content = re.sub(r'&#39;', "'", code_content)
            
            # Create placeholder
            placeholder = f"__CODE_BLOCK_{len(placeholder_map)}__"
            placeholder_map[placeholder] = code_content.strip()
            
            # Replace in HTML
            processed_html = processed_html[:match.start()] + placeholder + processed_html[match.end():]
    
    # Now clean all remaining HTML tags
    # Remove script and style content entirely
    processed_html = re.sub(r'<script[^>]*>.*?</script>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)
    processed_html = re.sub(r'<style[^>]*>.*?</style>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)
    
    # Convert common HTML elements to readable text
    # Headers
    processed_html = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\n\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Paragraphs
    processed_html = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Line breaks
    processed_html = re.sub(r'<br\s*/?>', '\n', processed_html, flags=re.IGNORECASE)
    # List items
    processed_html = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove all remaining HTML tags
    processed_html = re.sub(r'<[^>]+>', '', processed_html)
    
    # Clean up HTML entities
    processed_html = re.sub(r'&nbsp;', ' ', processed_html)
    processed_html = re.sub(r'&lt;', '<', processed_html)
    processed_html = re.sub(r'&gt;', '>', processed_html)
    processed_html = re.sub(r'&amp;', '&', processed_html)
    processed_html = re.sub(r'&quot;', '"', processed_html)
    processed_html = re.sub(r'&#39;', "'", processed_html)
    processed_html = re.sub(r'&#x27;', "'", processed_html)
    
    # Restore code blocks
    for placeholder, code_content in placeholder_map.items():
        processed_html = processed_html.replace(placeholder, f"\n\n```\n{code_content}\n```\n\n")
    
    # Clean up excessive whitespace
    processed_html = re.sub(r'\n\s*\n\s*\n', '\n\n', processed_html)  # Max 2 consecutive newlines
    processed_html = re.sub(r'[ \t]+', ' ', processed_html)  # Multiple spaces to single space
    
    return processed_html.strip()


def extract_text_from_document(file_content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from various document formats.

    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # PDF files
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return extract_text_from_pdf(file_content)

        # Word documents
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file_content)

        # HTML files - clean tags and extract text
        elif content_type == "text/html" or filename.lower().endswith((".html", ".htm")):
            # Decode HTML and clean tags for RAG
            html_text = file_content.decode("utf-8", errors="ignore").strip()
            if not html_text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return _clean_html_to_text(html_text)

        # Text files (markdown, txt, etc.)
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt",
            ".md",
            ".markdown",
            ".rst",
        )):
            # Decode text and check if it has content
            text = file_content.decode("utf-8", errors="ignore").strip()
            if not text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return text

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except ValueError:
        # Re-raise ValueError with original message for unsupported formats
        raise
    except Exception as e:
        logfire.error(
            "Document text extraction failed",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        # Re-raise with context, preserving original exception chain
        raise Exception(f"Failed to extract text from {filename}") from e


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using both PyPDF2 and pdfplumber for best results.

    Args:
        file_content: Raw PDF bytes

    Returns:
        Extracted text content
    """
    if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise Exception(
            "No PDF processing libraries available. Please install pdfplumber and PyPDF2."
        )

    text_content = []

    # First try with pdfplumber (better for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logfire.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue

            # If pdfplumber got good results, use them
            if text_content and len("\n".join(text_content).strip()) > 100:
                combined_text = "\n\n".join(text_content)
                logger.info(f"🔍 PDF DEBUG: Extracted {len(text_content)} pages, total length: {len(combined_text)}")
                logger.info(f"🔍 PDF DEBUG: First 500 chars: {repr(combined_text[:500])}")
                
                # Check for backticks before and after processing
                backtick_count_before = combined_text.count("```")
                logger.info(f"🔍 PDF DEBUG: Backticks found before processing: {backtick_count_before}")
                
                processed_text = _preserve_code_blocks_across_pages(combined_text)
                backtick_count_after = processed_text.count("```")
                logger.info(f"🔍 PDF DEBUG: Backticks found after processing: {backtick_count_after}")
                
                if backtick_count_after > 0:
                    logger.info(f"🔍 PDF DEBUG: Sample after processing: {repr(processed_text[:1000])}")
                
                return processed_text

        except Exception as e:
            logfire.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")

    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logfire.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            if text_content:
                combined_text = "\n\n".join(text_content)
                return _preserve_code_blocks_across_pages(combined_text)
            else:
                raise ValueError(
                    "No text extracted from PDF: file may be empty, images-only, "
                    "or scanned document without OCR"
                )

        except Exception as e:
            raise Exception("PyPDF2 failed to extract text") from e

    # If we get here, no libraries worked
    raise Exception("Failed to extract text from PDF - no working PDF libraries available")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word documents (.docx).

    Args:
        file_content: Raw DOCX bytes

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")

    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        if not text_content:
            raise ValueError("No text content found in document")

        return "\n\n".join(text_content)

    except Exception as e:
        raise Exception("Failed to extract text from Word document") from e
